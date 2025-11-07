"""
Script to update the Executive Report Word document from the updated Markdown file.
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def markdown_to_word(md_file, docx_file):
    """Convert updated Markdown to Word document with formatting."""
    
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split content into lines for processing
    lines = md_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip markdown code blocks
        if line.startswith('```'):
            i += 1
            continue
        
        # Main title (# Title)
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(192, 39, 45)  # Red color from logo
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.space_after = Pt(12)
        
        # Level 2 headings (## Heading)
        elif line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(192, 39, 45)
            p.space_before = Pt(12)
            p.space_after = Pt(6)
        
        # Level 3 headings (### Heading)
        elif line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.space_before = Pt(8)
            p.space_after = Pt(4)
        
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run('_' * 50)
            run.font.color.rgb = RGBColor(200, 200, 200)
        
        # Bold text with checkmark (✓ or ✅)
        elif line.startswith('✓') or line.startswith('✅'):
            p = doc.add_paragraph(line, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
        
        # Warning symbol (⚠)
        elif line.startswith('⚠'):
            p = doc.add_paragraph(line, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
        
        # Bullet points (- or *)
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Handle bold text
            if '**' in text:
                p = doc.add_paragraph(style='List Bullet')
                parts = text.split('**')
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    if idx % 2 == 1:  # Odd indices are bold
                        run.font.bold = True
            else:
                doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            # Handle bold text
            if '**' in text:
                p = doc.add_paragraph(style='List Number')
                parts = text.split('**')
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    if idx % 2 == 1:
                        run.font.bold = True
            else:
                doc.add_paragraph(text, style='List Number')
        
        # Table rows (|)
        elif line.startswith('|') and '|' in line[1:]:
            # Collect table rows
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # Step back one
            
            # Skip separator rows
            table_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:]+\|', l)]
            
            if table_lines:
                # Parse table
                rows = []
                for tline in table_lines:
                    cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                    rows.append(cells)
                
                if rows:
                    # Create table
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    # Fill table
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data.replace('**', '')
                            
                            # Bold header row
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                cell._element.get_or_add_tcPr().append(
                                    cell._element._new_tcPr()
                                )
                    
                    doc.add_paragraph()  # Space after table
        
        # Block quote (>)
        elif line.startswith('>'):
            text = line[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run(text.replace('*', ''))
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Regular paragraph with bold/italic handling
        elif line and not line.startswith('#'):
            if '**' in line or '*' in line:
                p = doc.add_paragraph()
                # Simple bold/italic parsing
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.font.italic = True
                    elif part:
                        p.add_run(part)
            elif line:
                doc.add_paragraph(line)
        
        # Empty line
        else:
            if i > 0 and lines[i-1].strip():  # Don't add multiple blank lines
                doc.add_paragraph()
        
        i += 1
    
    # Save the document
    doc.save(docx_file)
    print(f"✅ Successfully updated {docx_file}")

if __name__ == "__main__":
    md_file = r"c:\Users\karim\genai_clinicaltrials\reports\PROJECT_REPORT_EXECUTIVE.md"
    docx_file = r"c:\Users\karim\genai_clinicaltrials\reports\PROJECT_REPORT_EXECUTIVE.docx"
    
    try:
        markdown_to_word(md_file, docx_file)
        print("\n📄 Word document has been updated with all changes from the Markdown file.")
        print("   - Title updated to 'ClinicalIQ – Clinical Protocol Intelligence & Q&A'")
        print("   - Version 2.0 with multi-turn extraction updates")
        print("   - All performance metrics and test results updated")
        print("   - SME feedback and roadmap revised")
    except Exception as e:
        print(f"❌ Error: {e}")
        print(f"   Error type: {type(e).__name__}")
        import traceback
        traceback.print_exc()
        print("\nNote: python-docx is already installed. Check if the file paths are correct.")
