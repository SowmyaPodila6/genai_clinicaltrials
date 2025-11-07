"""
Generate Executive Report (Word only - concise 3-4 pages)
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_executive_word_report():
    """Generate concise Word document from executive markdown"""
    try:
        # Read markdown file
        with open('PROJECT_REPORT_EXECUTIVE.md', 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Set narrow margins for more content per page
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Title page
        title = doc.add_heading('Clinical Trial Analysis System', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Executive Report')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 128)
        
        doc.add_paragraph()
        
        info = doc.add_paragraph('Date: November 6, 2025  |  Version: 1.0')
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()
        doc.add_paragraph('_' * 80)
        
        # Process markdown content
        lines = md_content.split('\n')
        in_code_block = False
        code_content = []
        in_table = False
        table_rows = []
        skip_header = True
        
        for line in lines:
            # Skip title section (already added)
            if skip_header:
                if line.strip() == '---' and 'Executive Summary' in '\n'.join(lines[lines.index(line)+1:lines.index(line)+10]):
                    skip_header = False
                continue
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    if code_content:
                        p = doc.add_paragraph('\n'.join(code_content))
                        for run in p.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(8)
                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # Handle horizontal rules
            if line.strip() == '---':
                doc.add_paragraph()
                continue
            
            # Handle headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            
            # Handle tables
            elif '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []
                # Skip separator rows
                if line.strip().replace('|', '').replace('-', '').strip():
                    table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
            elif in_table and (line.strip() == '' or not line.strip().startswith('|')):
                # End of table
                if table_rows and len(table_rows) > 1:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for i, row_data in enumerate(table_rows):
                        row = table.rows[i]
                        for j, cell_data in enumerate(row_data):
                            row.cells[j].text = cell_data
                            # Format cells
                            for paragraph in row.cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    if i == 0:  # Header row
                                        run.font.bold = True
                    
                    doc.add_paragraph()
                in_table = False
                table_rows = []
                
                # Process the non-table line
                if line.strip() and not line.startswith('#'):
                    # Fall through to regular paragraph handling below
                    pass
                else:
                    continue
            
            # Skip if we're in a table
            if in_table:
                continue
            
            # Handle bullet points
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                text = text.replace('**', '').replace('`', '')
                p = doc.add_paragraph(text, style='List Bullet')
                for run in p.runs:
                    run.font.size = Pt(10)
            
            # Handle numbered lists
            elif line.strip() and len(line.strip()) > 2 and line.strip()[0].isdigit() and '. ' in line[:5]:
                text = line.strip().split('. ', 1)[1] if '. ' in line else line.strip()
                text = text.replace('**', '').replace('`', '')
                p = doc.add_paragraph(text, style='List Number')
                for run in p.runs:
                    run.font.size = Pt(10)
            
            # Handle blockquotes
            elif line.startswith('>'):
                text = line[1:].strip()
                p = doc.add_paragraph(text)
                p.style = 'Quote'
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.italic = True
            
            # Handle regular paragraphs
            elif line.strip() and not line.startswith('#'):
                text = line.strip()
                
                # Skip if it's part of skipped header
                if any(x in text for x in ['Project:', 'Date:', 'Version:', 'Repository:']) and text.startswith('**'):
                    continue
                
                # Handle checkmarks and X marks
                text = text.replace('✓', '[OK]').replace('✗', '[X]').replace('⚠', '[!]')
                
                if text:
                    p = doc.add_paragraph()
                    
                    # Parse bold formatting
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            # Normal text - handle inline code
                            if '`' in part:
                                code_parts = part.split('`')
                                for j, code_part in enumerate(code_parts):
                                    if j % 2 == 0:
                                        if code_part:
                                            run = p.add_run(code_part)
                                            run.font.size = Pt(10)
                                    else:
                                        run = p.add_run(code_part)
                                        run.font.name = 'Courier New'
                                        run.font.size = Pt(9)
                            else:
                                if part:
                                    run = p.add_run(part)
                                    run.font.size = Pt(10)
                        else:
                            # Bold text
                            if part:
                                run = p.add_run(part)
                                run.bold = True
                                run.font.size = Pt(10)
        
        # Save document
        output_path = 'PROJECT_REPORT_EXECUTIVE.docx'
        doc.save(output_path)
        print(f"[OK] Executive Word document generated: {output_path}")
        print(f"     Size: {os.path.getsize(output_path) / 1024:.1f} KB")
        print(f"     Pages: Approximately 3-4 pages")
        return output_path
        
    except Exception as e:
        print(f"[X] Error generating executive Word document: {e}")
        import traceback
        traceback.print_exc()
        return None


if __name__ == "__main__":
    print("=" * 60)
    print("Clinical Trial Analysis System - Executive Report Generator")
    print("=" * 60)
    print()
    
    if not os.path.exists('PROJECT_REPORT_EXECUTIVE.md'):
        print("[X] Error: PROJECT_REPORT_EXECUTIVE.md not found!")
    else:
        print("Generating concise executive report (3-4 pages)...")
        print()
        generate_executive_word_report()
        print()
        print("Markdown source: PROJECT_REPORT_EXECUTIVE.md")
