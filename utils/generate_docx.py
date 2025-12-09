#!/usr/bin/env python3
"""
Simple script to convert PROJECT_REPORT.md to DOCX format
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    # Remove bold and italic formatting
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    # Remove links
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    return text

def markdown_to_docx(md_file_path, docx_file_path):
    """Convert markdown file to Word document"""
    
    # Create new document
    doc = Document()
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
        
        # Skip code blocks
        if line.startswith('```'):
            continue
        
        # Headers
        if line.startswith('#'):
            level = line.count('#')
            text = clean_markdown_text(line.lstrip('#').strip())
            
            if level == 1:
                heading = doc.add_heading(text, 1)
            elif level == 2:
                heading = doc.add_heading(text, 2)
            elif level == 3:
                heading = doc.add_heading(text, 3)
            else:
                heading = doc.add_heading(text, 4)
        
        # Regular paragraphs
        elif not line.startswith('---') and not line.startswith('|'):
            cleaned_text = clean_markdown_text(line)
            if cleaned_text.strip():
                doc.add_paragraph(cleaned_text)
    
    # Save document
    doc.save(docx_file_path)
    print(f"✅ DOCX report generated: {docx_file_path}")

def main():
    """Main function"""
    # File paths
    reports_dir = Path("reports")
    md_file = reports_dir / "PROJECT_REPORT.md"
    docx_file = reports_dir / "PROJECT_REPORT.docx"
    
    # Check if markdown file exists
    if not md_file.exists():
        print(f"❌ Markdown file not found: {md_file}")
        return
    
    # Create reports directory if it doesn't exist
    reports_dir.mkdir(exist_ok=True)
    
    # Convert to DOCX
    markdown_to_docx(md_file, docx_file)
    print("✅ Report generation completed successfully!")

if __name__ == "__main__":
    main()