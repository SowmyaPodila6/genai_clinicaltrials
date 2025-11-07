"""
Generate Word and PDF versions of the project report
Requires: python-docx, markdown, pdfkit/weasyprint
"""

import os
from pathlib import Path

def generate_word_report():
    """Generate Word document from markdown"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import markdown
        from bs4 import BeautifulSoup
        
        # Read markdown file
        with open('PROJECT_REPORT.md', 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML for parsing
        html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        soup = BeautifulSoup(html, 'html.parser')
        
        # Create Word document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Clinical Trial Analysis System - Project Report"
        doc.core_properties.author = "Clinical Trial Analysis Team"
        doc.core_properties.comments = "Comprehensive project report including architecture, results, and roadmap"
        
        # Add title page
        title = doc.add_heading('Clinical Trial Analysis System', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Project Report')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(18)
        subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 128)
        
        doc.add_paragraph()
        
        date_para = doc.add_paragraph('Date: November 6, 2025')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        version_para = doc.add_paragraph('Version: 1.0')
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Process markdown content line by line
        lines = md_content.split('\n')
        in_code_block = False
        code_content = []
        table_started = False
        table_rows = []
        
        for line in lines:
            # Skip title page content (already added)
            if line.startswith('# Clinical Trial Analysis System - Project Report'):
                continue
            if 'Project Name:' in line or 'Date:' in line or 'Version:' in line or 'Repository:' in line:
                if line.startswith('**'):
                    continue
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    if code_content:
                        p = doc.add_paragraph('\n'.join(code_content))
                        # Use monospace font instead of 'Code' style which might not exist
                        for run in p.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        code_content = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # Handle horizontal rules
            if line.strip() == '---':
                doc.add_paragraph('_' * 80)
                continue
            
            # Handle headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            
            # Handle tables (markdown tables)
            elif '|' in line and line.strip().startswith('|'):
                # Simple table handling
                if not table_started:
                    table_started = True
                    table_rows = []
                table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
            elif table_started and line.strip() == '':
                # End of table
                if table_rows:
                    # Create Word table
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for i, row_data in enumerate(table_rows):
                        row = table.rows[i]
                        for j, cell_data in enumerate(row_data):
                            row.cells[j].text = cell_data
                            # Bold header row
                            if i == 0:
                                row.cells[j].paragraphs[0].runs[0].font.bold = True
                    
                    doc.add_paragraph()
                table_started = False
                table_rows = []
            
            # Handle bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                # Remove markdown formatting
                text = text.replace('**', '').replace('`', '')
                doc.add_paragraph(text, style='List Bullet')
            
            # Handle numbered lists
            elif line.strip() and line.strip()[0].isdigit() and '. ' in line[:5]:
                text = line.strip().split('. ', 1)[1]
                text = text.replace('**', '').replace('`', '')
                doc.add_paragraph(text, style='List Number')
            
            # Handle blockquotes
            elif line.startswith('>'):
                text = line[1:].strip()
                p = doc.add_paragraph(text)
                p.style = 'Quote'
            
            # Handle regular paragraphs
            elif line.strip() and not table_started:
                # Process inline formatting
                text = line.strip()
                
                # Skip separator lines
                if text.startswith('┌') or text.startswith('│') or text.startswith('└') or text.startswith('─'):
                    p = doc.add_paragraph(text)
                    p.runs[0].font.name = 'Courier New'
                    continue
                
                if text:
                    p = doc.add_paragraph()
                    
                    # Simple bold/italic handling
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            # Normal text
                            if '`' in part:
                                code_parts = part.split('`')
                                for j, code_part in enumerate(code_parts):
                                    if j % 2 == 0:
                                        if code_part:
                                            p.add_run(code_part)
                                    else:
                                        run = p.add_run(code_part)
                                        run.font.name = 'Courier New'
                                        run.font.size = Pt(9)
                            else:
                                if part:
                                    p.add_run(part)
                        else:
                            # Bold text
                            if part:
                                run = p.add_run(part)
                                run.bold = True
        
        # Save Word document
        output_path = 'PROJECT_REPORT.docx'
        doc.save(output_path)
        print(f"[OK] Word document generated: {output_path}")
        return output_path
        
    except ImportError as e:
        print(f"X Missing required library: {e}")
        print("Install with: pip install python-docx markdown beautifulsoup4")
        return None
    except Exception as e:
        print(f"X Error generating Word document: {e}")
        return None


def generate_pdf_report():
    """Generate PDF document from markdown"""
    try:
        import markdown
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        
        # Read markdown file
        with open('PROJECT_REPORT.md', 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_content, 
            extensions=['tables', 'fenced_code', 'nl2br']
        )
        
        # Create styled HTML
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                    @bottom-right {{
                        content: "Page " counter(page) " of " counter(pages);
                        font-size: 9pt;
                        color: #666;
                    }}
                }}
                
                body {{
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                }}
                
                h1 {{
                    color: #003366;
                    border-bottom: 3px solid #003366;
                    padding-bottom: 10px;
                    page-break-before: always;
                    margin-top: 40px;
                }}
                
                h1:first-of-type {{
                    page-break-before: avoid;
                    margin-top: 0;
                    text-align: center;
                    font-size: 28pt;
                }}
                
                h2 {{
                    color: #0066cc;
                    border-bottom: 2px solid #0066cc;
                    padding-bottom: 8px;
                    margin-top: 30px;
                }}
                
                h3 {{
                    color: #0066cc;
                    margin-top: 25px;
                }}
                
                h4 {{
                    color: #333;
                    margin-top: 20px;
                }}
                
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 20px 0;
                    font-size: 10pt;
                }}
                
                th {{
                    background-color: #003366;
                    color: white;
                    padding: 10px;
                    text-align: left;
                    font-weight: bold;
                }}
                
                td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                }}
                
                tr:nth-child(even) {{
                    background-color: #f2f2f2;
                }}
                
                code {{
                    background-color: #f4f4f4;
                    padding: 2px 6px;
                    border-radius: 3px;
                    font-family: 'Courier New', monospace;
                    font-size: 9pt;
                }}
                
                pre {{
                    background-color: #f4f4f4;
                    padding: 15px;
                    border-radius: 5px;
                    border-left: 4px solid #003366;
                    overflow-x: auto;
                    font-family: 'Courier New', monospace;
                    font-size: 8pt;
                    line-height: 1.4;
                }}
                
                blockquote {{
                    border-left: 4px solid #0066cc;
                    padding-left: 20px;
                    margin-left: 0;
                    color: #555;
                    font-style: italic;
                }}
                
                ul, ol {{
                    margin: 15px 0;
                    padding-left: 30px;
                }}
                
                li {{
                    margin: 8px 0;
                }}
                
                strong {{
                    color: #003366;
                }}
                
                hr {{
                    border: none;
                    border-top: 2px solid #ddd;
                    margin: 30px 0;
                }}
                
                .page-break {{
                    page-break-after: always;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Generate PDF
        output_path = 'PROJECT_REPORT.pdf'
        font_config = FontConfiguration()
        
        HTML(string=styled_html).write_pdf(
            output_path,
            font_config=font_config
        )
        
        print(f"[OK] PDF document generated: {output_path}")
        return output_path
        
    except ImportError as e:
        print(f"X Missing required library for PDF: {e}")
        print("Install with: pip install weasyprint markdown")
        print("\nAlternatively, you can use the Word document and export to PDF manually.")
        return None
    except Exception as e:
        print(f"X Error generating PDF document: {e}")
        print("Note: WeasyPrint requires additional system dependencies.")
        print("Alternative: Use the Word document (.docx) and export to PDF manually.")
        return None


def main():
    """Generate both Word and PDF reports"""
    print("=" * 60)
    print("Clinical Trial Analysis System - Report Generator")
    print("=" * 60)
    print()
    
    # Check if markdown file exists
    if not os.path.exists('PROJECT_REPORT.md'):
        print("X Error: PROJECT_REPORT.md not found!")
        return
    
    print("Generating project reports...")
    print()
    
    # Generate Word document
    print("[1/2] Generating Word document...")
    word_path = generate_word_report()
    print()
    
    # Generate PDF document
    print("[2/2] Generating PDF document...")
    pdf_path = generate_pdf_report()
    print()
    
    # Summary
    print("=" * 60)
    print("Summary:")
    print("=" * 60)
    
    if word_path:
        print(f"[OK] Word Report: {word_path}")
        print(f"  Size: {os.path.getsize(word_path) / 1024:.1f} KB")
    else:
        print("X Word Report: Failed")
    
    if pdf_path:
        print(f"[OK] PDF Report: {pdf_path}")
        print(f"  Size: {os.path.getsize(pdf_path) / 1024:.1f} KB")
    else:
        print("X PDF Report: Failed (you can export Word to PDF manually)")
    
    print()
    print("Markdown source available at: PROJECT_REPORT.md")
    print()
    
    # Installation instructions if libraries missing
    if not word_path or not pdf_path:
        print("=" * 60)
        print("Installation Instructions:")
        print("=" * 60)
        print()
        print("To install required dependencies:")
        print("  pip install python-docx markdown beautifulsoup4 weasyprint")
        print()
        print("Note: WeasyPrint may require additional system libraries:")
        print("  - Windows: Install GTK+ for Windows")
        print("  - macOS: brew install cairo pango gdk-pixbuf libffi")
        print("  - Linux: apt-get install python3-dev python3-pip python3-cffi")
        print()


if __name__ == "__main__":
    main()
